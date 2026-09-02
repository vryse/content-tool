"""Parse reference documents while retaining structure needed for style analysis."""

from __future__ import annotations

import asyncio
import logging
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document

from app.models import ParsedArticle, ReferenceRecord, Section
from app.utils.storage import cached_copy

LOGGER = logging.getLogger(__name__)
HEADING_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)
MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")
MARKDOWN_TABLE_SEPARATOR_RE = re.compile(
    r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$", re.MULTILINE
)
MERMAID_FENCE_RE = re.compile(r"```mermaid[^\n]*\n(.*?)```", re.IGNORECASE | re.DOTALL)
MERMAID_FLOW_RE = re.compile(
    r"^\s*(?:flowchart|graph)\s+(?:TB|TD|BT|RL|LR)\b", re.IGNORECASE | re.MULTILINE
)


def _has_mermaid_flowchart(text: str) -> bool:
    return any(MERMAID_FLOW_RE.search(block) for block in MERMAID_FENCE_RE.findall(text))


def _section(heading: str, level: int, paragraphs: list[str], bullets: int, numbered: int) -> Section:
    return Section(
        heading=heading,
        level=level,
        paragraphs=paragraphs,
        word_count=sum(len(paragraph.split()) for paragraph in paragraphs),
        bullet_count=bullets,
        numbered_count=numbered,
    )


def load_docx_article(path: str | Path, company: str) -> ParsedArticle:
    """Parse paragraphs only; images and tables are intentionally harmless extras."""
    path = Path(path)
    document = Document(path)
    title = path.stem
    sections: list[Section] = []
    heading = "__intro__"
    level = 0
    paragraphs: list[str] = []
    bullets = numbered = 0
    saw_heading = False

    def flush() -> None:
        nonlocal paragraphs, bullets, numbered
        if paragraphs or heading != "__intro__":
            sections.append(_section(heading, level, paragraphs, bullets, numbered))
        paragraphs, bullets, numbered = [], 0, 0

    for paragraph in document.paragraphs:
        text = " ".join(paragraph.text.split())
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        if style.casefold() == "title":
            title = text
            continue
        match = HEADING_RE.match(style)
        if match:
            flush()
            heading, level = text, int(match.group(1))
            saw_heading = True
            continue
        paragraphs.append(text)
        style_lower = style.casefold()
        bullets += int(style_lower.startswith("list bullet"))
        numbered += int(style_lower.startswith("list number"))

    flush()
    if not saw_heading:
        LOGGER.warning("%s has no headings; using a single synthetic section", path.name)
        all_paragraphs = [p for section in sections for p in section.paragraphs]
        sections = [
            _section(
                "__intro__",
                0,
                all_paragraphs,
                sum(section.bullet_count for section in sections),
                sum(section.numbered_count for section in sections),
            )
        ]
    full_text = "\n\n".join(p for section in sections for p in section.paragraphs)
    return ParsedArticle(
        filename=path.name,
        title=title,
        sections=sections,
        full_text=full_text,
        company=company,
        has_table=bool(document.tables),
    )


def load_markdown_article(path: str | Path, company: str) -> ParsedArticle:
    """Keep Markdown headings and paragraphs useful to the style profiler."""
    path = Path(path)
    text = path.read_text(encoding="utf-8", errors="replace")
    title = path.stem.replace("-", " ").replace("_", " ").strip() or path.stem
    sections: list[Section] = []
    heading, level, paragraphs = "__intro__", 0, []

    def flush() -> None:
        nonlocal paragraphs
        if paragraphs or heading != "__intro__":
            sections.append(_section(heading, level, paragraphs, 0, 0))
        paragraphs = []

    for line in text.splitlines():
        match = MARKDOWN_HEADING_RE.match(line.strip())
        if match:
            flush()
            heading, level = match.group(2), len(match.group(1))
            if level == 1:
                title = heading
            continue
        cleaned = line.strip()
        if cleaned:
            paragraphs.append(cleaned)
    flush()
    if not sections:
        raise ValueError("contains no readable Markdown text")
    return ParsedArticle(
        filename=path.name,
        title=title,
        sections=sections,
        full_text="\n\n".join(p for section in sections for p in section.paragraphs),
        company=company,
        has_table=bool(MARKDOWN_TABLE_SEPARATOR_RE.search(text)),
        has_flowchart=_has_mermaid_flowchart(text),
    )


def load_doc_article(path: str | Path, company: str) -> ParsedArticle:
    """Extract a legacy Word document with the OS converter, then profile its text.

    ``python-docx`` deliberately cannot read binary .doc files. macOS ships
    ``textutil``; deployments without it get an actionable error rather than a
    corrupted pseudo-document.
    """
    path = Path(path)
    with tempfile.TemporaryDirectory() as directory:
        output = Path(directory) / "reference.txt"
        try:
            subprocess.run(
                ["textutil", "-convert", "txt", "-output", str(output), str(path)],
                check=True,
                capture_output=True,
                text=True,
                timeout=30,
            )
        except FileNotFoundError as error:
            raise ValueError("legacy .doc import requires macOS textutil or a .docx conversion") from error
        except subprocess.CalledProcessError as error:
            raise ValueError(error.stderr.strip() or "could not convert legacy .doc") from error
        text = output.read_text(encoding="utf-8", errors="replace")
    scratch = path.with_suffix(".converted.md")
    try:
        scratch.write_text(f"# {path.stem}\n\n{text}", encoding="utf-8")
        return load_markdown_article(scratch, company).model_copy(update={"filename": path.name})
    finally:
        scratch.unlink(missing_ok=True)


def load_article(path: str | Path, company: str) -> ParsedArticle:
    """Parse the supported reference formats into the common article contract."""
    path = Path(path)
    suffix = path.suffix.casefold()
    if suffix == ".docx":
        return load_docx_article(path, company)
    if suffix in {".md", ".markdown"}:
        return load_markdown_article(path, company)
    if suffix == ".doc":
        return load_doc_article(path, company)
    raise ValueError(f"Unsupported reference format: {suffix or 'unknown'}")


def load_local(reference_dir: str | Path, company: str) -> list[ParsedArticle]:
    """Parse a directory of .docx files; used by the one-off R2 seeding script."""
    return [load_article(path, company) for path in sorted(Path(reference_dir).glob("*.docx"))]


async def load_reference(record: ReferenceRecord) -> ParsedArticle:
    """Parse one stored reference, reading its bytes through the local R2 cache."""
    path = await cached_copy(record.key, record.content_hash)
    article = await asyncio.to_thread(load_article, path, record.company)
    # The cache filename is a content hash, so the original name has to be
    # restored here or every skeleton the retriever returns would be unlabelled.
    return article.model_copy(update={"filename": record.filename})


async def load_references(records: list[ReferenceRecord]) -> list[ParsedArticle]:
    """Parse a selection concurrently, skipping documents that will not open.

    A single unreadable upload must not take down profile building for the rest
    of the corpus; the failure is already recorded against the row and shown in
    the ingest UI.
    """
    ordered = sorted(records, key=lambda record: record.key)
    results = await asyncio.gather(
        *(load_reference(record) for record in ordered), return_exceptions=True
    )
    articles: list[ParsedArticle] = []
    for record, result in zip(ordered, results):
        if isinstance(result, BaseException):
            LOGGER.warning("Skipping unreadable reference %s: %s", record.key, result)
            continue
        articles.append(result)
    return articles
